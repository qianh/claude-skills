#!/usr/bin/env python3
"""
试卷Word文档生成器
使用 python-docx 库生成专业的试卷文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


class ExamPaperGenerator:
    """试卷生成器类"""

    def __init__(self):
        self.doc = Document()
        self._set_default_font()

    def _set_default_font(self):
        """设置默认字体为中文友好字体"""
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(10.5)

    def set_header(self, school="", subject="", duration=120, total_score=100,
                   semester="", class_info=""):
        """
        设置试卷头部信息

        Args:
            school: 学校名称
            subject: 科目名称
            duration: 考试时长（分钟）
            total_score: 总分
            semester: 学期（如 "2024学年第一学期"）
            class_info: 班级信息
        """
        # 学校名称（居中，大字号）
        if school:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(school)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 试卷标题（居中，大字号加粗）
        title = f"{semester} {subject}试卷" if semester else f"{subject}试卷"
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 空行
        self.doc.add_paragraph()

        # 考试信息（居中）
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_text = f"考试时间：{duration}分钟    总分：{total_score}分"
        if class_info:
            info_text = f"{class_info}    " + info_text
        run = p.add_run(info_text)
        run.font.size = Pt(10.5)

        # 学生信息填写栏（居中）
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("姓名：__________    学号：__________    班级：__________")
        run.font.size = Pt(10.5)

        # 分隔线
        self.doc.add_paragraph("_" * 70)

    def add_section(self, section_title, questions=None):
        """
        添加一个大题部分

        Args:
            section_title: 大题标题（如 "一、选择题（共30分）"）
            questions: 题目列表（可选）
        """
        # 大题标题（加粗）
        p = self.doc.add_paragraph()
        run = p.add_run(section_title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加题目
        if questions:
            for q in questions:
                self._add_question(q)

        # 空行
        self.doc.add_paragraph()

    def _add_question(self, question_dict):
        """
        添加单个题目

        Args:
            question_dict: 包含题目信息的字典
                {
                    'number': 1,
                    'stem': '题干',
                    'options': ['A. 选项A', 'B. 选项B', ...],  # 可选
                    'points': 5,  # 可选
                    'image_path': 'path/to/image.png'  # 可选
                }
        """
        # 题号和题干
        stem_text = f"{question_dict['number']}. "
        if 'points' in question_dict:
            stem_text += f"（{question_dict['points']}分）"
        stem_text += question_dict['stem']

        p = self.doc.add_paragraph(stem_text)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)

        # 如果有图片，插入图片
        if 'image_path' in question_dict and os.path.exists(question_dict['image_path']):
            self.add_image(question_dict['image_path'], width_cm=12)

        # 如果有选项（选择题）
        if 'options' in question_dict:
            for option in question_dict['options']:
                p = self.doc.add_paragraph(option)
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_after = Pt(3)

        # 答题空间（如果需要）
        if question_dict.get('answer_space'):
            lines = question_dict.get('answer_lines', 3)
            for _ in range(lines):
                p = self.doc.add_paragraph("_" * 60)
                p.paragraph_format.left_indent = Cm(1.5)

    def add_image(self, image_path, width_cm=10):
        """
        插入图片

        Args:
            image_path: 图片路径
            width_cm: 图片宽度（厘米）
        """
        if os.path.exists(image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Cm(width_cm))
            p.paragraph_format.space_after = Pt(12)
        else:
            print(f"警告：图片文件不存在：{image_path}")

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()

    def add_answer_section(self, answers_dict):
        """
        添加参考答案部分

        Args:
            answers_dict: 答案字典，格式示例：
                {
                    '一、选择题': [
                        {'number': 1, 'answer': 'A', 'explanation': '解析...'},
                        ...
                    ],
                    '二、简答题': [
                        {'number': 1, 'answer': '答案内容', 'points_breakdown': '评分标准'},
                        ...
                    ]
                }
        """
        # 新开一页
        self.add_page_break()

        # 答案标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("参考答案及评分标准")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 各部分答案
        for section_title, answers in answers_dict.items():
            # 部分标题
            p = self.doc.add_paragraph()
            run = p.add_run(section_title)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # 各题答案
            for ans in answers:
                ans_text = f"{ans['number']}. {ans['answer']}"
                p = self.doc.add_paragraph(ans_text)
                p.paragraph_format.left_indent = Cm(0.5)

                # 解析或评分标准
                if 'explanation' in ans and ans['explanation']:
                    exp_p = self.doc.add_paragraph(f"   解析：{ans['explanation']}")
                    exp_p.paragraph_format.left_indent = Cm(1.0)
                    exp_p.paragraph_format.space_after = Pt(6)

                if 'points_breakdown' in ans and ans['points_breakdown']:
                    pb_p = self.doc.add_paragraph(f"   评分标准：{ans['points_breakdown']}")
                    pb_p.paragraph_format.left_indent = Cm(1.0)
                    pb_p.paragraph_format.space_after = Pt(6)

            # 空行
            self.doc.add_paragraph()

    def add_table(self, data, headers=None):
        """
        添加表格

        Args:
            data: 二维列表
            headers: 表头列表（可选）
        """
        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else 0

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'

        # 添加表头
        if headers:
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            data_start_row = 1
        else:
            data_start_row = 0

        # 添加数据
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                table.rows[data_start_row + i].cells[j].text = str(cell_data)

        self.doc.add_paragraph()

    def save(self, filename):
        """
        保存文档

        Args:
            filename: 文件名（应以.docx结尾）
        """
        if not filename.endswith('.docx'):
            filename += '.docx'

        self.doc.save(filename)
        print(f"试卷已保存：{filename}")
        return filename


# 使用示例
if __name__ == "__main__":
    # 创建生成器
    generator = ExamPaperGenerator()

    # 设置头部
    generator.set_header(
        school="XX中学",
        subject="数学",
        semester="2024学年第一学期",
        duration=120,
        total_score=100
    )

    # 添加选择题部分
    choice_questions = [
        {
            'number': 1,
            'stem': '下列函数中，在定义域内单调递增的是',
            'options': [
                'A. y = -x²',
                'B. y = 1/x',
                'C. y = 2^x',
                'D. y = |x|'
            ],
            'points': 5
        },
        {
            'number': 2,
            'stem': '集合 A = {1, 2, 3}，B = {2, 3, 4}，则 A ∩ B =',
            'options': [
                'A. {1}',
                'B. {2, 3}',
                'C. {1, 2, 3, 4}',
                'D. ∅'
            ],
            'points': 5
        }
    ]
    generator.add_section("一、选择题（每题5分，共10分）", choice_questions)

    # 添加简答题部分
    short_answer = [
        {
            'number': 1,
            'stem': '已知函数 f(x) = 2x + 1，求 f(f(x)) 的表达式。',
            'points': 8,
            'answer_space': True,
            'answer_lines': 5
        }
    ]
    generator.add_section("二、简答题（共8分）", short_answer)

    # 添加答案
    answers = {
        '一、选择题（每题5分，共10分）': [
            {'number': 1, 'answer': 'C', 'explanation': '指数函数 y = 2^x 在定义域 R 上单调递增'},
            {'number': 2, 'answer': 'B', 'explanation': '交集是两个集合共有的元素，即 {2, 3}'}
        ],
        '二、简答题（共8分）': [
            {
                'number': 1,
                'answer': 'f(f(x)) = f(2x+1) = 2(2x+1)+1 = 4x+3',
                'points_breakdown': '正确代入得4分，计算正确得4分'
            }
        ]
    }
    generator.add_answer_section(answers)

    # 保存
    generator.save("示例试卷.docx")
